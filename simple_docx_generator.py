"""
Générateur simple qui remplit un template Word en remplaçant directement le texte
"""
from docx import Document
import os
import subprocess
import re


class SimpleDocxGenerator:
    """Génère un CPV en remplissant directement le template DOCX"""

    def __init__(self, template_path):
        self.template_path = template_path

    def generate(self, output_path, extracted_data, form_data):
        """
        Génère le DOCX rempli

        Args:
            output_path: Chemin de sortie
            extracted_data: Données extraites
            form_data: Données du formulaire
        """
        # Charger le template
        doc = Document(self.template_path)

        # Créer le mapping de remplacement
        replacements = self._build_replacements(extracted_data, form_data)

        # Remplacer dans les paragraphes
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)

        # Remplacer dans les tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)

        # Sauvegarder le DOCX
        doc.save(output_path)
        print(f"✅ DOCX généré: {output_path}")

        # Essayer de convertir en PDF
        pdf_path = output_path.replace('.docx', '.pdf')
        if self._convert_to_pdf(output_path, pdf_path):
            print(f"✅ PDF généré: {pdf_path}")
            return pdf_path
        else:
            print(f"📄 DOCX généré (conversion PDF non disponible)")
            return output_path

    def _clean_text(self, text):
        """Nettoie le texte pour le rendre compatible XML"""
        if text is None or text == '':
            return ''

        text = str(text)
        # Supprimer les caractères de contrôle et NULL bytes
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
        return text

    def _replace_in_paragraph(self, paragraph, replacements):
        """Remplace le texte dans un paragraphe"""
        for old, new in replacements.items():
            if old in paragraph.text:
                # Nettoyer la nouvelle valeur
                clean_new = self._clean_text(new)
                # Remplacer en préservant le style
                for run in paragraph.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, clean_new)

    def _convert_to_pdf(self, docx_path, pdf_path):
        """Tente de convertir le DOCX en PDF avec LibreOffice"""
        try:
            # Essayer LibreOffice sur macOS
            libreoffice_path = '/Applications/LibreOffice.app/Contents/MacOS/soffice'
            if os.path.exists(libreoffice_path):
                output_dir = os.path.dirname(pdf_path)
                subprocess.run([
                    libreoffice_path,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', output_dir,
                    docx_path
                ], check=True, capture_output=True, timeout=30)
                return True
        except:
            pass

        return False

    def _build_replacements(self, extracted_data, form_data):
        """Construit le dictionnaire de remplacement"""

        segment = extracted_data.get('segment', 'C4').upper()

        replacements = {
            # Nom du client (apparaît plusieurs fois)
            'mallet': extracted_data.get('raison_sociale', ''),
            'Nom du client : mallet': f"Nom du client : {extracted_data.get('raison_sociale', '')}",

            # Informations entreprise
            'NOM CLIENT': extracted_data.get('raison_sociale', ''),
            'Type société': extracted_data.get('forme_juridique', ''),
            'Montant Capital Social': form_data.get('capital_social', extracted_data.get('capital_social', '')),
            'Adresse Siège Social Client': extracted_data.get('adresse_siege', ''),
            'Ville RCS': extracted_data.get('ville_rcs', ''),
            'Numéro RCS': extracted_data.get('siren', ''),
            'Nom Prénom signataire': extracted_data.get('signataire_nom', ''),
            'Fonction signataire': form_data.get('fonction_signataire', extracted_data.get('fonction_signataire', '')),
            'Nom Client': extracted_data.get('raison_sociale', ''),

            # Dates
            '01/01/2026ut_Fourniture': extracted_data.get('date_debut_livraison', '01/01/2026'),
            'Date_Fin_ Fourniture': extracted_data.get('date_fin_livraison', '31/12/2026'),

            # Garantie
            'XXX €': f"{form_data.get('garantie_montant', '0')} €",

            # Contact
            'A compléter par le Client': extracted_data.get('email', ''),

            # Nom signataire
            'Nom Signataire': extracted_data.get('signataire_nom', ''),
            'Fonction Signataire': form_data.get('fonction_signataire', extracted_data.get('fonction_signataire', '')),

            # PRM et volume
            '… MWh': f"{extracted_data.get('volume_total', '')} MWh",

            # IBAN/BIC
            '_ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _': form_data.get('iban', ''),
            '_ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _ _': form_data.get('bic', ''),
        }

        # Prix abonnement et capacité (communs à tous)
        if form_data.get('prix_abonnement'):
            replacements['€/an'] = f"{form_data.get('prix_abonnement')} €/an"

        if form_data.get('prix_capacite'):
            replacements['Prix Capacité'] = form_data.get('prix_capacite', '')

        # Ajouter consommations et prix selon segment
        if segment == 'C2':
            # 5 postes: PTE, HPH, HCH, HPE, HCE
            postes = ['pte', 'hph', 'hch', 'hpe', 'hce']
            for poste in postes:
                conso_key = f'conso_{poste}'
                prix_key = f'prix_{poste}'

                if form_data.get(conso_key):
                    replacements[poste.upper()] = form_data.get(conso_key, '')

                if form_data.get(prix_key):
                    replacements[f'Prix {poste.upper()}'] = form_data.get(prix_key, '')

        elif segment == 'C4':
            # 4 postes: HPH, HCH, HPE, HCE
            postes = ['hph', 'hch', 'hpe', 'hce']
            for poste in postes:
                conso_key = f'conso_{poste}'
                prix_key = f'prix_{poste}'

                if form_data.get(conso_key):
                    replacements[poste.upper()] = form_data.get(conso_key, '')

                if form_data.get(prix_key):
                    replacements[f'Prix {poste.upper()}'] = form_data.get(prix_key, '')

        elif segment == 'C5':
            c5_option = form_data.get('c5_option', 'base')

            if c5_option == 'base':
                # Option BASE
                if form_data.get('conso_base'):
                    replacements['BASE'] = form_data.get('conso_base', '')
                if form_data.get('prix_base'):
                    replacements['Prix BASE'] = form_data.get('prix_base', '')
            else:
                # Option HP/HC
                if form_data.get('conso_hp'):
                    replacements['HP'] = form_data.get('conso_hp', '')
                if form_data.get('conso_hc'):
                    replacements['HC'] = form_data.get('conso_hc', '')
                if form_data.get('prix_hp'):
                    replacements['Prix HP'] = form_data.get('prix_hp', '')
                if form_data.get('prix_hc'):
                    replacements['Prix HC'] = form_data.get('prix_hc', '')

        # Garanties d'Origine
        if form_data.get('go_souhaite') == 'oui':
            if form_data.get('go_percentage'):
                replacements['Prix GO'] = f"{form_data.get('go_percentage')}% couvert"

        # CEE
        if form_data.get('cee'):
            replacements['Client soumis aux CEE'] = 'Soumis' if form_data.get('cee') == 'soumis' else 'Non Soumis'

        return replacements
