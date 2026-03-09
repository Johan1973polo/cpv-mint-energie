#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur DOCX CPV MINT 2026
Remplit le template Word avec les données du formulaire
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import json
import os
from lxml import etree

# CONFIRMATION DE CHARGEMENT DU NOUVEAU CODE (2026-02-28 13:10)
print("=" * 80)
print("✅ MODULE docx_generator_2026.py RECHARGÉ (VERSION AVEC FIX TEXTE VERTICAL)")
print("=" * 80)


class CPVGenerator2026:
    """Générateur de CPV 2026 basé sur template Word"""

    # Mapping FTA longues → codes courts pour éviter retour à la ligne
    FTA_SHORT = {
        'BT > 36 kVA': 'BTSUP36',
        'BT ≤ 36 kVA': 'BTINF36',
        'BT <= 36 kVA': 'BTINF36',
        'HTA': 'HTA',
        'HTLU': 'HTLU',
        'HTCU': 'HTCU',
        'MU4': 'MU4',
        'CU4': 'CU4',
        'LU4': 'LU4',
        'BTINF': 'BTINF'
    }

    def __init__(self, template_path='template_cpv_2026.docx'):
        """
        Args:
            template_path: Chemin vers le template Word
        """
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template introuvable: {template_path}")

        self.template_path = template_path

    def _raccourcir_fta(self, fta):
        """Raccourcit les codes FTA trop longs pour éviter retour à la ligne"""
        if not fta:
            return ''
        return self.FTA_SHORT.get(str(fta).strip(), str(fta).strip())


    def _set_cell_text(self, cell, text, font_size=8, bold=False, center=True):
        """Écrit du texte dans une cellule avec le formatage MINT (8pt, centré)"""
        # IMPORTANT: Supprimer TOUS les paragraphes existants et en créer un nouveau
        # (sinon le formatage du template persiste et affiche le texte verticalement)
        for paragraph in cell.paragraphs:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)

        # DÉSACTIVER LE WORD WRAP dans la cellule (évite le retour à la ligne automatique)
        # Cela empêche Word de couper le texte sur plusieurs lignes dans des colonnes étroites
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Ajouter <w:noWrap/> si pas déjà présent
        noWrap = tcPr.find(qn('w:noWrap'))
        if noWrap is None:
            noWrap = etree.SubElement(tcPr, qn('w:noWrap'))

        # Créer un nouveau paragraphe propre
        paragraph = cell.add_paragraph()
        if center:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Ajouter le texte avec la bonne taille
        run = paragraph.add_run(str(text) if text else '')
        run.font.size = Pt(font_size)
        run.font.bold = bold

    def generate(self, output_path, extracted_data, form_data, totaux_par_segment=None):
        """
        Génère le CPV en remplissant le template

        Args:
            output_path: Chemin du fichier de sortie
            extracted_data: Données extraites des PDFs
            form_data: Données du formulaire
            totaux_par_segment: Dictionnaire avec les sites groupés par segment (C2, C4, C5)

        Returns:
            str: Chemin du fichier généré
        """
        print(f"\n📝 Génération CPV 2026...")
        print(f"   Template: {self.template_path}")
        print(f"   Output: {output_path}")

        # Charger le template
        doc = Document(self.template_path)

        # Fusionner les données
        all_data = {**extracted_data, **form_data}

        # Remplir les paragraphes
        self._fill_paragraphs(doc, all_data)

        # Remplir les tableaux avec les sites groupés par segment
        self._fill_tables(doc, all_data, form_data, totaux_par_segment)

        # Sauvegarder
        doc.save(output_path)
        print(f"✅ CPV généré: {output_path}")

        return output_path

    def _fill_paragraphs(self, doc, data):
        """Remplit les paragraphes du document en gérant les runs fragmentés"""
        print("\n📄 Remplissage des paragraphes...")

        # Tous les placeholders à remplacer
        forme_jur = data.get('forme_juridique', 'SAS')
        capital = data.get('capital_social', '')

        # Gérer le capital social : retirer le € car le template contient déjà " €"
        if capital:
            capital = str(capital).replace(' €', '').replace('€', '').strip()

        replacements = {
            'NOM CLIENT': data.get('raison_sociale', ''),
            'SIREN': data.get('siren', ''),
            'Montant Capital Social': capital,
            'Adresse Siège Social': data.get('adresse_siege', ''),
            'Ville RCS': data.get('ville_rcs', ''),
            'Prénom NOM signataire': data.get('nom_signataire', data.get('nom_gerant', '')),
            'Fonction signataire': data.get('fonction_signataire', 'Gérant'),
            'Forme Juridique': forme_jur,
            'Forme juridique': forme_jur,  # Variante avec minuscule
            'Forme juridique,': forme_jur  # Variante avec virgule
        }

        # Fonction helper pour remplacer dans un paragraphe (gère les runs fragmentés)
        def replace_in_paragraph(paragraph, old_text, new_text):
            """Remplace old_text par new_text dans un paragraphe, même si fragmenté sur plusieurs runs"""
            if not new_text:
                return False

            full_text = paragraph.text
            if old_text in full_text:
                # Vider tous les runs sauf le premier
                for i, run in enumerate(paragraph.runs):
                    if i == 0:
                        run.text = full_text.replace(old_text, str(new_text))
                    else:
                        run.text = ""
                print(f"   ✓ Remplacé '{old_text}' → '{str(new_text)[:40]}...'")
                return True
            return False

        # Appliquer les remplacements à tous les paragraphes
        for para in doc.paragraphs:
            for placeholder, value in replacements.items():
                replace_in_paragraph(para, placeholder, value)

        # Gestion spéciale pour Entrepreneur individuel (pas de capital social)
        if forme_jur == 'Entrepreneur individuel':
            print("\n📝 Entrepreneur individuel détecté - suppression de la ligne 'au capital de'")
            # Chercher et supprimer le paragraphe contenant "au capital de"
            paragraphs_to_remove = []
            for para in doc.paragraphs:
                if 'au capital de' in para.text.lower() or 'au capital social de' in para.text.lower():
                    paragraphs_to_remove.append(para)
                    print(f"   → Paragraphe à supprimer: {para.text[:60]}...")

            # Supprimer les paragraphes identifiés
            for para in paragraphs_to_remove:
                # Vider complètement le paragraphe
                for run in para.runs:
                    run.text = ""
                print(f"   ✓ Paragraphe vidé")

        # Gérer les cases à cocher (checkboxes) pour GO et CEE
        print("\n☑️  Gestion des checkboxes...")
        self._fill_checkboxes(doc, data)

    def _fill_checkboxes(self, doc, data):
        """Remplit les cases à cocher pour GO et CEE dans le document

        IMPORTANT: Le template contient déjà des symboles ☐ qu'il faut REMPLACER.
        Cette méthode les remplace par [X] (coché) ou [ ] (non coché).
        """
        # Récupérer les valeurs
        go_value = str(data.get('go_souhaite', '0'))
        cee_status = data.get('cee_status', 'non_soumis')

        print(f"\n☑️  Gestion des checkboxes...")
        print(f"   GO souhaité: {go_value}")
        print(f"   Statut CEE: {cee_status}")

        # DEBUG: Écrire le XML dans un fichier
        debug_file = "/Users/strategyglobal/Desktop/cpv mint/cpv_app/debug_xml.txt"
        with open(debug_file, 'w', encoding='utf-8') as f:
            f.write("=" * 80 + "\n")
            f.write("DEBUG: Recherche des paragraphes contenant 'Soumis'\n")
            f.write("=" * 80 + "\n\n")

            # Chercher dans les paragraphes principaux
            for i, para in enumerate(doc.paragraphs):
                if 'Soumis' in para.text:
                    f.write(f"\n--- Paragraphe doc {i} ---\n")
                    f.write(f"TEXT: {repr(para.text)}\n")
                    f.write(f"XML:\n{etree.tostring(para._element, pretty_print=True, encoding='unicode')}\n")
                    f.write("-" * 80 + "\n")

            # Chercher dans les tableaux
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        for para_idx, para in enumerate(cell.paragraphs):
                            if 'Soumis' in para.text:
                                f.write(f"\n--- Table {table_idx}, Row {row_idx}, Cell {cell_idx}, Para {para_idx} ---\n")
                                f.write(f"TEXT: {repr(para.text)}\n")
                                f.write(f"XML:\n{etree.tostring(para._element, pretty_print=True, encoding='unicode')}\n")
                                f.write("-" * 80 + "\n")

            f.write("=" * 80 + "\n")

        print(f"   🐛 Debug XML écrit dans: {debug_file}")

        # Définir les options CEE et GO avec leurs patterns
        # ORDRE IMPORTANT : du plus spécifique au plus générique (éviter "Soumis" matche avant "Non Soumis")
        # Utiliser des textes courts pour matcher même si le template a un texte tronqué
        cee_checks = [
            ('Non Soumis (tous les sites)', cee_status == 'non_soumis'),
            ('Mixte (selon secteur', cee_status == 'mixte'),  # Raccourci pour matcher même si tronqué
            ('Soumis (tous les sites)', cee_status == 'soumis'),
        ]

        go_checks = [
            ('Non souhaité', go_value == '0'),
            ('Souhaité', go_value != '0'),
            ('100%', go_value == '100'),
            ('50%', go_value == '50'),
            ('25%', go_value == '25'),
        ]

        def cocher_case(para, texte_recherche, doit_cocher):
            """Coche ou décoche une case en supprimant le CONTENT CONTROL CHECKBOX Word"""
            clean = para.text.strip()
            # Nettoyer pour comparaison (retirer ☐, □, espaces, tabs, unicode checkboxes)
            clean_compare = clean.lstrip('☐□▢\u2610\u2612\u25A1\u25A0 \t')

            # Vérifier si le texte correspond
            if clean_compare.startswith(texte_recherche):
                # SUPPRIMER les Content Controls <w:sdt> (checkboxes Word 2010+)
                nsmap = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                    'w14': 'http://schemas.microsoft.com/office/word/2010/wordml'
                }

                # Trouver tous les <w:sdt> dans le paragraphe
                for sdt in para._element.findall('.//w:sdt', nsmap):
                    # Vérifier si c'est une checkbox
                    checkbox = sdt.find('.//w14:checkbox', nsmap)
                    if checkbox is not None:
                        # SUPPRIMER TOUT le Content Control
                        para._element.remove(sdt)

                new_text = "[X] " + texte_recherche if doit_cocher else "[ ] " + texte_recherche

                # VIDER ABSOLUMENT TOUS LES RUNS et réécrire le PREMIER
                for i, run in enumerate(para.runs):
                    if i == 0:
                        # Réécrire le premier run avec le texte complet
                        run.text = new_text
                        run.bold = True if doit_cocher else False
                    else:
                        # Vider complètement les autres runs
                        run.text = ""

                return True
            return False

        # Parcourir tous les paragraphes
        for i, para in enumerate(doc.paragraphs):
            # Traiter les checkboxes CEE
            for texte, doit_cocher in cee_checks:
                if cocher_case(para, texte, doit_cocher):
                    if doit_cocher:
                        print(f"   ✓ P{i}: CEE '{texte}' coché")
                    break

            # Traiter les checkboxes GO
            for texte, doit_cocher in go_checks:
                if cocher_case(para, texte, doit_cocher):
                    if doit_cocher:
                        print(f"   ✓ P{i}: GO '{texte}' coché")
                    break

        # Parcourir aussi les tableaux (au cas où il y aurait des checkboxes dans les cellules)
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        # Traiter les checkboxes CEE dans tableaux
                        for texte, doit_cocher in cee_checks:
                            if cocher_case(para, texte, doit_cocher):
                                break

                        # Traiter les checkboxes GO dans tableaux
                        for texte, doit_cocher in go_checks:
                            if cocher_case(para, texte, doit_cocher):
                                break

        print(f"   ✅ Checkboxes mis à jour")

    def _fill_tables(self, doc, all_data, form_data, totaux_par_segment=None):
        """Remplit tous les tableaux du document"""
        print("\n📊 Remplissage des tableaux...")

        if len(doc.tables) == 0:
            print("   ⚠️  Aucun tableau trouvé dans le template")
            return

        # Tableau 0 : Périmètre contractuel - MULTI-SEGMENT
        if len(doc.tables) > 0:
            self._fill_perimetre(doc.tables[0], all_data, form_data, totaux_par_segment)

        # Tableau 1 : Facturation
        if len(doc.tables) > 1:
            self._fill_facturation(doc.tables[1], all_data)

        # Tableaux 2/3/4 : Prix C2/C4/C5 - Remplir chaque segment avec les données du premier site
        if len(doc.tables) > 2:
            self._fill_prix_c2(doc.tables[2], form_data, totaux_par_segment)
        if len(doc.tables) > 3:
            self._fill_prix_c4(doc.tables[3], form_data, totaux_par_segment)
        if len(doc.tables) > 4:
            self._fill_prix_c5(doc.tables[4], form_data, totaux_par_segment)

        # Tableaux 6/7/8 : Sites C2/C4/C5 - Remplir chaque segment avec ses sites
        if len(doc.tables) > 6:
            self._fill_sites_c2(doc.tables[6], form_data, totaux_par_segment)
        if len(doc.tables) > 7:
            self._fill_sites_c4(doc.tables[7], form_data, totaux_par_segment)
        if len(doc.tables) > 8:
            self._fill_sites_c5(doc.tables[8], form_data, totaux_par_segment)

        # Tableau 9 : Signatures
        if len(doc.tables) > 9:
            self._fill_signatures(doc.tables[9], all_data)

        # Tableau 10 : Interlocuteurs
        if len(doc.tables) > 10:
            self._fill_interlocuteurs(doc.tables[10], all_data)

        # Tableau 12 : SEPA
        if len(doc.tables) > 12:
            self._fill_sepa(doc.tables[12], all_data)

    def _fill_perimetre(self, table, all_data, form_data, totaux_par_segment=None):
        """Remplit le tableau 0 : Périmètre contractuel - MULTI-SEGMENT"""
        print("   → Tableau 0: Périmètre contractuel (Multi-segment)")

        try:
            # R2: dates début/fin (utiliser les dates du premier site ou dates globales)
            if len(table.rows) > 2:
                table.rows[2].cells[1].text = form_data.get('date_debut', '')
                table.rows[2].cells[2].text = form_data.get('date_fin', '')

            # Configuration des segments
            segments_config = {
                'C2': {'row': 4, 'flexibilite_key': 'flexibilite_c2'},
                'C4': {'row': 5, 'flexibilite_key': 'flexibilite_c4'},
                'C5': {'row': 6, 'flexibilite_key': 'flexibilite_c5'}
            }

            total_prm = 0
            total_volume = 0.0

            # Remplir TOUS les segments avec les données de totaux_par_segment
            for segment, config in segments_config.items():
                row_idx = config['row']
                if len(table.rows) > row_idx:
                    # Vérifier si ce segment a des sites
                    if totaux_par_segment and segment in totaux_par_segment:
                        segment_data = totaux_par_segment[segment]
                        nb_prm = segment_data['nb_prm']
                        volume = segment_data['volume_total']
                        flexibilite = form_data.get(config['flexibilite_key'], 'Non')

                        # Remplir la ligne
                        table.rows[row_idx].cells[1].text = str(nb_prm)
                        table.rows[row_idx].cells[2].text = f"{volume:.2f}"
                        table.rows[row_idx].cells[3].text = flexibilite

                        total_prm += nb_prm
                        total_volume += volume

                        print(f"      {segment}: {nb_prm} PRM, {volume:.2f} MWh")
                    else:
                        # Aucun site pour ce segment
                        table.rows[row_idx].cells[1].text = "—"
                        table.rows[row_idx].cells[2].text = "—"
                        table.rows[row_idx].cells[3].text = "—"
                        print(f"      {segment}: Aucun site")

            # R7: Total (somme de tous les segments)
            if len(table.rows) > 7:
                table.rows[7].cells[1].text = str(total_prm)
                table.rows[7].cells[2].text = f"{total_volume:.2f}"

            print(f"      Total global: {total_prm} PRM, {total_volume:.2f} MWh")

        except Exception as e:
            print(f"      ⚠️  Erreur: {e}")
            import traceback
            traceback.print_exc()

    def _fill_facturation(self, table, data):
        """Remplit le tableau 1 : Facturation"""
        print("   → Tableau 1: Facturation")

        try:
            # R1: Adresse facturation
            if len(table.rows) > 1:
                table.rows[1].cells[1].text = data.get('adresse_facturation',
                                                       data.get('adresse_siege', ''))

            # R2: Email facturation
            if len(table.rows) > 2:
                table.rows[2].cells[1].text = data.get('email_facturation',
                                                       data.get('email', ''))

            # R8: Garantie paiement
            if len(table.rows) > 8:
                table.rows[8].cells[1].text = data.get('garantie_paiement', 'Non')

        except Exception as e:
            print(f"      ⚠️  Erreur: {e}")

    def _fill_prix_c2(self, table, form_data, totaux_par_segment=None):
        """Remplit le tableau 2 : Prix C2 - Utilise le premier site C2"""
        # Vérifier si il y a des sites C2
        has_c2 = totaux_par_segment and 'C2' in totaux_par_segment and totaux_par_segment['C2']['sites']
        print(f"   → Tableau 2: Prix C2 {'(REMPLI)' if has_c2 else '(IGNORÉ - Aucun site C2)'}")

        if not has_c2:
            return

        try:
            # Récupérer les prix depuis le PREMIER site C2
            premier_site_c2 = totaux_par_segment['C2']['sites'][0]
            prix_p0_data = premier_site_c2.get('prix_p0_data', {})

            # DEBUG: Afficher ce qui est reçu
            print(f"      🔍 DEBUG C2:")
            print(f"         Site C2: PRM={premier_site_c2.get('prm', 'N/A')}")
            print(f"         prix_p0_data type: {type(prix_p0_data)}")
            print(f"         prix_p0_data: {prix_p0_data}")

            if not prix_p0_data:
                print(f"      ⚠️  Aucune donnée de prix pour le premier site C2")
                return

            prix_finaux = prix_p0_data.get('prix_finaux', {})
            print(f"         prix_finaux: {prix_finaux}")
            print(f"         Clés disponibles: {list(prix_finaux.keys())}")

            # R1: Prix électricité
            if len(table.rows) > 1:
                table.rows[1].cells[1].text = f"{prix_finaux.get('prix_pte', 0):.2f}"
                table.rows[1].cells[2].text = f"{prix_finaux.get('prix_hph', 0):.2f}"
                table.rows[1].cells[3].text = f"{prix_finaux.get('prix_hch', 0):.2f}"
                table.rows[1].cells[4].text = f"{prix_finaux.get('prix_hpe', 0):.2f}"
                table.rows[1].cells[5].text = f"{prix_finaux.get('prix_hce', 0):.2f}"

            # R2: Coefficient α (C2 = POINTE uniquement)
            if len(table.rows) > 2:
                coef_alpha = prix_p0_data.get('coefficient_alpha', 0.7)
                alpha_str = f"{float(coef_alpha):g}"

                table.rows[2].cells[1].text = alpha_str  # POINTE/PTE
                table.rows[2].cells[2].text = "-"        # HPH
                table.rows[2].cells[3].text = "-"        # HCH
                table.rows[2].cells[4].text = "-"        # HPE
                table.rows[2].cells[5].text = "-"        # HCE

                print(f"      ✓ Coefficient α={alpha_str} positionné sur POINTE uniquement")

            print(f"      ✓ Prix C2 remplis depuis site PRM={premier_site_c2.get('prm', 'N/A')}")

        except Exception as e:
            print(f"      ⚠️  Erreur: {e}")
            import traceback
            traceback.print_exc()

    def _fill_prix_c4(self, table, form_data, totaux_par_segment=None):
        """Remplit le tableau 3 : Prix C4 - Utilise le premier site C4"""
        # Vérifier si il y a des sites C4
        has_c4 = totaux_par_segment and 'C4' in totaux_par_segment and totaux_par_segment['C4']['sites']
        print(f"   → Tableau 3: Prix C4 {'(REMPLI)' if has_c4 else '(IGNORÉ - Aucun site C4)'}")

        if not has_c4:
            return

        try:
            # Récupérer les prix depuis le PREMIER site C4
            premier_site_c4 = totaux_par_segment['C4']['sites'][0]
            prix_p0_data = premier_site_c4.get('prix_p0_data', {})

            if not prix_p0_data:
                print(f"      ⚠️  Aucune donnée de prix pour le premier site C4")
                return

            prix_finaux = prix_p0_data.get('prix_finaux', {})

            # R1: Prix électricité (5 colonnes : Segment C4 | Pointe/HPH | HCH | HPE | HCE)
            if len(table.rows) > 1 and len(table.rows[1].cells) >= 5:
                hph = prix_finaux.get('prix_hph', 0)
                hch = prix_finaux.get('prix_hch', 0)
                hpe = prix_finaux.get('prix_hpe', 0)
                hce = prix_finaux.get('prix_hce', 0)

                table.rows[1].cells[1].text = f"{hph:.2f}"
                table.rows[1].cells[2].text = f"{hch:.2f}"
                table.rows[1].cells[3].text = f"{hpe:.2f}"
                table.rows[1].cells[4].text = f"{hce:.2f}"

            # R2: Coefficient α (C4 = HPH uniquement)
            if len(table.rows) > 2 and len(table.rows[2].cells) >= 5:
                coef_alpha = prix_p0_data.get('coefficient_alpha', 0.7)
                alpha_str = f"{float(coef_alpha):g}"

                table.rows[2].cells[1].text = alpha_str  # HPH (Pointe/HPH)
                table.rows[2].cells[2].text = "-"        # HCH
                table.rows[2].cells[3].text = "-"        # HPE
                table.rows[2].cells[4].text = "-"        # HCE

                print(f"      ✓ Coefficient α={alpha_str} positionné sur HPH uniquement")

            print(f"      ✓ Prix C4 remplis depuis site PRM={premier_site_c4.get('prm', 'N/A')}")

        except Exception as e:
            print(f"      ⚠️  Erreur: {e}")
            import traceback
            traceback.print_exc()

    def _fill_prix_c5(self, table, form_data, totaux_par_segment=None):
        """Remplit le tableau 4 : Prix C5 - Utilise le premier site C5"""
        # Vérifier si il y a des sites C5
        has_c5 = totaux_par_segment and 'C5' in totaux_par_segment and totaux_par_segment['C5']['sites']
        print(f"   → Tableau 4: Prix C5 {'(REMPLI)' if has_c5 else '(IGNORÉ - Aucun site C5)'}")

        if not has_c5:
            return

        try:
            # Récupérer les prix depuis le PREMIER site C5
            premier_site_c5 = totaux_par_segment['C5']['sites'][0]
            prix_p0_data = premier_site_c5.get('prix_p0_data', {})

            print(f"      🔍 DEBUG C5:")
            print(f"         premier_site_c5 PRM: {premier_site_c5.get('prm', 'N/A')}")
            print(f"         prix_p0_data type: {type(prix_p0_data)}")
            print(f"         prix_p0_data: {prix_p0_data}")

            if not prix_p0_data:
                print(f"      ⚠️  Aucune donnée de prix pour le premier site C5")
                return

            prix_finaux = prix_p0_data.get('prix_finaux', {})
            print(f"         prix_finaux type: {type(prix_finaux)}")
            print(f"         prix_finaux: {prix_finaux}")
            print(f"         Clés dans prix_finaux: {list(prix_finaux.keys()) if isinstance(prix_finaux, dict) else 'N/A'}")

            # Détecter les types de calendriers utilisés parmi TOUS les sites C5
            sites_c5 = totaux_par_segment['C5']['sites']
            has_base = False
            has_hphc = False
            has_4cadrans = False

            for site in sites_c5:
                type_cal = site.get('type_calendrier', '').upper()
                if 'BASE' in type_cal:
                    has_base = True
                elif 'HP' in type_cal and 'HC' in type_cal:
                    has_hphc = True
                elif '4' in type_cal or 'CADRAN' in type_cal:
                    has_4cadrans = True

            print(f"      🔍 Calendriers détectés sur {len(sites_c5)} site(s): BASE={has_base}, HP-HC={has_hphc}, 4 cadrans={has_4cadrans}")

            # R1: Prix électricité
            if len(table.rows) > 1:
                base = prix_finaux.get('prix_base', 0)
                hp = prix_finaux.get('prix_hp', 0)
                hc = prix_finaux.get('prix_hc', 0)
                hph = prix_finaux.get('prix_hpsh', prix_finaux.get('prix_hph', 0))
                hch = prix_finaux.get('prix_hcsh', prix_finaux.get('prix_hch', 0))
                hpe = prix_finaux.get('prix_hpsb', prix_finaux.get('prix_hpe', 0))
                hce = prix_finaux.get('prix_hcsb', prix_finaux.get('prix_hce', 0))

                table.rows[1].cells[1].text = f"{base:.2f}" if has_base else "-"
                table.rows[1].cells[2].text = f"{hp:.2f}" if has_hphc else "-"
                table.rows[1].cells[3].text = f"{hc:.2f}" if has_hphc else "-"
                table.rows[1].cells[4].text = f"{hph:.2f}" if has_4cadrans else "-"
                table.rows[1].cells[5].text = f"{hch:.2f}" if has_4cadrans else "-"
                table.rows[1].cells[6].text = f"{hpe:.2f}" if has_4cadrans else "-"
                table.rows[1].cells[7].text = f"{hce:.2f}" if has_4cadrans else "-"

            # R2: Coefficient α
            if len(table.rows) > 2:
                coef_alpha = prix_p0_data.get('coefficient_alpha', 0.7)
                alpha_str = f"{float(coef_alpha):g}"
                nb_types_calendrier = sum([has_base, has_hphc, has_4cadrans])

                alpha_cells = {'BASE': '-', 'HP': '-', 'HC': '-', 'HPH': '-', 'HCH': '-', 'HPE': '-', 'HCE': '-'}

                if nb_types_calendrier == 1:
                    if has_base:
                        alpha_cells['BASE'] = alpha_str
                        positions = "BASE uniquement"
                    elif has_hphc:
                        alpha_cells['HP'] = alpha_str
                        positions = "HP uniquement"
                    elif has_4cadrans:
                        alpha_cells['HPH'] = alpha_str
                        positions = "HPH uniquement"
                else:
                    active_positions = []
                    if has_base:
                        alpha_cells['BASE'] = alpha_str
                        active_positions.append('BASE')
                    if has_hphc:
                        alpha_cells['HP'] = alpha_str
                        active_positions.append('HP')
                    if has_4cadrans:
                        alpha_cells['HPH'] = alpha_str
                        active_positions.append('HPH')
                    positions = " + ".join(active_positions)

                table.rows[2].cells[1].text = alpha_cells['BASE']
                table.rows[2].cells[2].text = alpha_cells['HP']
                table.rows[2].cells[3].text = alpha_cells['HC']
                table.rows[2].cells[4].text = alpha_cells['HPH']
                table.rows[2].cells[5].text = alpha_cells['HCH']
                table.rows[2].cells[6].text = alpha_cells['HPE']
                table.rows[2].cells[7].text = alpha_cells['HCE']

                print(f"      ✓ Coefficient α={alpha_str} positionné sur: {positions}")

            print(f"      ✓ Prix C5 remplis depuis site PRM={premier_site_c5.get('prm', 'N/A')}")

        except Exception as e:
            print(f"      ⚠️  Erreur: {e}")
            import traceback
            traceback.print_exc()

    def _fill_sites_c2(self, table, form_data, totaux_par_segment=None):
        """Remplit le tableau 6 : Sites C2"""
        has_c2 = totaux_par_segment and 'C2' in totaux_par_segment and totaux_par_segment['C2']['sites']
        print(f"   → Tableau 6: Sites C2 {'(REMPLI)' if has_c2 else '(IGNORÉ - Aucun site C2)'}")

        # Supprimer toutes les lignes d'exemple (après l'en-tête = row 0)
        while len(table.rows) > 1:
            table._element.remove(table.rows[1]._element)

        if not has_c2:
            print(f"      ✓ Lignes d'exemple supprimées")
            return

        try:
            sites_c2 = totaux_par_segment['C2']['sites']

            # Récupérer la valeur CEE GLOBALE
            cee_status_global = form_data.get('cee_status', 'non_soumis')
            cee_display_global = {'non_soumis': 'Non', 'soumis': 'Oui', 'mixte': 'Mixte'}.get(cee_status_global, cee_status_global)

            # Ajouter une ligne pour chaque site C2
            for site in sites_c2:
                new_row = table.add_row()
                if len(new_row.cells) >= 13:
                    self._set_cell_text(new_row.cells[0], site['prm'])
                    self._set_cell_text(new_row.cells[1], site['siret'])
                    self._set_cell_text(new_row.cells[2], site['naf'])
                    self._set_cell_text(new_row.cells[3], site['adresse'])
                    self._set_cell_text(new_row.cells[4], cee_display_global)
                    self._set_cell_text(new_row.cells[5], 'C2')
                    self._set_cell_text(new_row.cells[6], self._raccourcir_fta(site['fta']))
                    self._set_cell_text(new_row.cells[7], site['date_debut'])
                    self._set_cell_text(new_row.cells[8], site['date_fin'])
                    # Colonnes 9-12: Puissance pour chaque poste
                    puissance = site.get('puissance', '')
                    self._set_cell_text(new_row.cells[9], puissance)   # PTE
                    self._set_cell_text(new_row.cells[10], puissance)  # HPH
                    self._set_cell_text(new_row.cells[11], puissance)  # HCH
                    self._set_cell_text(new_row.cells[12], puissance)  # HPE

            print(f"      ✓ {len(sites_c2)} site(s) C2 ajouté(s)")

        except Exception as e:
            print(f"      ⚠️  Erreur: {e}")
            import traceback
            traceback.print_exc()

    def _fill_sites_c4(self, table, form_data, totaux_par_segment=None):
        """Remplit le tableau 7 : Sites C4"""
        has_c4 = totaux_par_segment and 'C4' in totaux_par_segment and totaux_par_segment['C4']['sites']
        print(f"   → Tableau 7: Sites C4 {'(REMPLI)' if has_c4 else '(IGNORÉ - Aucun site C4)'}")

        # Supprimer toutes les lignes d'exemple (après l'en-tête = row 0)
        while len(table.rows) > 1:
            table._element.remove(table.rows[1]._element)

        if not has_c4:
            print(f"      ✓ Lignes d'exemple supprimées")
            return

        try:
            sites_c4 = totaux_par_segment['C4']['sites']

            # Récupérer la valeur CEE GLOBALE
            cee_status_global = form_data.get('cee_status', 'non_soumis')
            cee_display_global = {'non_soumis': 'Non', 'soumis': 'Oui', 'mixte': 'Mixte'}.get(cee_status_global, cee_status_global)

            # Ajouter une ligne pour chaque site C4
            for site in sites_c4:
                new_row = table.add_row()
                if len(new_row.cells) >= 13:
                    self._set_cell_text(new_row.cells[0], site['prm'])
                    self._set_cell_text(new_row.cells[1], site['siret'])
                    self._set_cell_text(new_row.cells[2], site['naf'])
                    self._set_cell_text(new_row.cells[3], site['adresse'])
                    self._set_cell_text(new_row.cells[4], cee_display_global)
                    self._set_cell_text(new_row.cells[5], 'C4')
                    self._set_cell_text(new_row.cells[6], self._raccourcir_fta(site['fta']))
                    self._set_cell_text(new_row.cells[7], site['date_debut'])
                    self._set_cell_text(new_row.cells[8], site['date_fin'])
                    # Colonnes 9-12: Puissance pour chaque poste
                    puissance = site.get('puissance', '')
                    self._set_cell_text(new_row.cells[9], puissance)   # Pointe/HPH
                    self._set_cell_text(new_row.cells[10], puissance)  # HCH
                    self._set_cell_text(new_row.cells[11], puissance)  # HPE
                    self._set_cell_text(new_row.cells[12], puissance)  # HCE

            print(f"      ✓ {len(sites_c4)} site(s) C4 ajouté(s)")

        except Exception as e:
            print(f"      ⚠️  Erreur: {e}")
            import traceback
            traceback.print_exc()

    def _fill_sites_c5(self, table, form_data, totaux_par_segment=None):
        """Remplit le tableau 8 : Sites C5"""
        has_c5 = totaux_par_segment and 'C5' in totaux_par_segment and totaux_par_segment['C5']['sites']
        print(f"   → Tableau 8: Sites C5 {'(REMPLI)' if has_c5 else '(IGNORÉ - Aucun site C5)'}")

        # Supprimer toutes les lignes d'exemple (après l'en-tête = row 0)
        while len(table.rows) > 1:
            table._element.remove(table.rows[1]._element)

        if not has_c5:
            print(f"      ✓ Lignes d'exemple supprimées")
            return

        try:
            sites_c5 = totaux_par_segment['C5']['sites']

            # Récupérer la valeur CEE GLOBALE
            cee_status_global = form_data.get('cee_status', 'non_soumis')
            cee_display_global = {'non_soumis': 'Non', 'soumis': 'Oui', 'mixte': 'Mixte'}.get(cee_status_global, cee_status_global)

            # Ajouter une ligne pour chaque site C5
            for site in sites_c5:
                # Formater le type de calendrier
                type_cal_display = site.get('type_calendrier', '').replace('4_cadrans', '4 cadrans')

                new_row = table.add_row()
                if len(new_row.cells) >= 11:
                    self._set_cell_text(new_row.cells[0], site['prm'])
                    self._set_cell_text(new_row.cells[1], site['siret'])
                    self._set_cell_text(new_row.cells[2], site['naf'])
                    self._set_cell_text(new_row.cells[3], site['adresse'])
                    self._set_cell_text(new_row.cells[4], cee_display_global)
                    self._set_cell_text(new_row.cells[5], 'C5')
                    self._set_cell_text(new_row.cells[6], self._raccourcir_fta(site['fta']))
                    self._set_cell_text(new_row.cells[7], site['date_debut'])
                    self._set_cell_text(new_row.cells[8], site['date_fin'])
                    self._set_cell_text(new_row.cells[9], type_cal_display)
                    self._set_cell_text(new_row.cells[10], site['puissance'])

            print(f"      ✓ {len(sites_c5)} site(s) C5 ajouté(s)")

        except Exception as e:
            print(f"      ⚠️  Erreur: {e}")
            import traceback
            traceback.print_exc()

    def _fill_signatures(self, table, data):
        """Remplit le tableau 9 : Signatures"""
        print("   → Tableau 9: Signatures")

        try:
            if len(table.rows) > 0:
                # Combiner nom + prénom si disponibles séparément
                nom = data.get('nom_signataire', data.get('nom_gerant', ''))
                prenom = data.get('prenom_signataire', data.get('prenom_gerant', ''))

                # Si prénom existe, combiner "NOM Prénom", sinon juste le nom
                if prenom:
                    nom_complet = f"{nom} {prenom}"
                else:
                    nom_complet = nom

                fonction = data.get('fonction_signataire', 'Gérant')

                # Utiliser _set_cell_text pour éviter la séparation en colonnes
                # Formatage gauche aligné, 10pt pour la lisibilité
                self._set_cell_text(
                    table.rows[0].cells[1],
                    f"{nom_complet}\n{fonction}",
                    font_size=10,
                    center=False
                )
                print(f"      ✓ Signataire: {nom_complet} ({fonction})")

        except Exception as e:
            print(f"      ⚠️  Erreur: {e}")

    def _fill_interlocuteurs(self, table, data):
        """Remplit le tableau 10 : Interlocuteurs - MAPPING CORRIGÉ"""
        print("   → Tableau 10: Interlocuteurs")

        try:
            # Mapping correct selon spécification utilisateur:
            # R0 col1 : Nom interlocuteur
            # R1 col1 : Adresse physique
            # R2 col1 : Email
            # R3 col1 : Téléphone fixe
            # R4 col1 : Téléphone mobile

            contacts = {
                0: data.get('contact_technique_nom', data.get('nom_signataire', data.get('nom_gerant', ''))),
                1: data.get('adresse_siege', ''),
                2: data.get('email', ''),
                3: data.get('telephone', ''),
                4: data.get('telephone_mobile', data.get('telephone', ''))
            }

            for row_idx, contact_value in contacts.items():
                if row_idx < len(table.rows) and contact_value:
                    table.rows[row_idx].cells[1].text = contact_value
                    print(f"      R{row_idx}: {contact_value[:40]}...")

        except Exception as e:
            print(f"      ⚠️  Erreur: {e}")

    def _fill_sepa(self, table, data):
        """Remplit le tableau 12 : SEPA"""
        print("   → Tableau 12: SEPA")

        try:
            # R1: Raison sociale + adresse - utiliser _set_cell_text pour éviter séparation
            if len(table.rows) > 1:
                raison_sociale = data.get('raison_sociale', '')
                adresse = data.get('adresse_siege', '')

                # Utiliser _set_cell_text pour garantir que tout reste dans cells[0]
                self._set_cell_text(
                    table.rows[1].cells[0],
                    f"{raison_sociale}\n{adresse}",
                    font_size=10,
                    center=False
                )
                print(f"      ✓ Raison sociale: {raison_sociale}")

            # R2: IBAN + BIC
            if len(table.rows) > 2:
                iban = data.get('iban', '')
                bic = data.get('bic', '')

                self._set_cell_text(
                    table.rows[2].cells[0],
                    f"IBAN: {iban}\nBIC: {bic}",
                    font_size=10,
                    center=False
                )
                print(f"      ✓ IBAN: {iban[:20]}...")

        except Exception as e:
            print(f"      ⚠️  Erreur: {e}")


if __name__ == '__main__':
    """Test du générateur"""
    import sys

    if len(sys.argv) < 2:
        print("Usage: python docx_generator_2026.py output.docx")
        sys.exit(1)

    # Données de test
    extracted_data = {
        'raison_sociale': 'TEST SAS',
        'siren': '123456789',
        'adresse_siege': '1 rue de Test, 75001 Paris',
        'nom_gerant': 'Jean DUPONT',
        'email': 'test@example.com'
    }

    form_data = {
        'date_debut': '01/03/2026',
        'date_fin': '28/02/2027',
        'site_count': '1',
        'prix_p0_data': json.dumps({
            'prix_finaux': {'pte': 100.5, 'hph': 120.0},
            'coefficient_alpha': 0.1234
        })
    }

    generator = CPVGenerator2026()
    output_path = sys.argv[1]
    generator.generate(output_path, extracted_data, form_data)
    print(f"\n✅ Test OK: {output_path}")
